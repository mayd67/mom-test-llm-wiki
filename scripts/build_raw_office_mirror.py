from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = ROOT / "raw"
WIKI_DIR = ROOT / "wiki"
TODAY = date.today().isoformat()

GROUP_OUTPUTS = {
    "产品资料": WIKI_DIR / "03_业务系统" / "MOM" / "原始资料镜像",
    "需求文档": WIKI_DIR / "03_业务系统" / "MOM" / "原始资料镜像",
    "接口资料": WIKI_DIR / "03_业务系统" / "MOM" / "原始资料镜像",
    "项目文档": WIKI_DIR / "03_业务系统" / "MOM" / "原始资料镜像",
    "截图附件": WIKI_DIR / "03_业务系统" / "MOM" / "原始资料镜像",
    "测试流程规范": WIKI_DIR / "02_测试标准&模板" / "原始资料镜像",
    "测试资料": WIKI_DIR / "02_测试标准&模板" / "原始资料镜像",
}

TARGET_EXTENSIONS = {".docx", ".xlsx"}
IGNORED_FILE_NAMES = {".gitkeep", ".gitignore"}
IGNORED_DIR_NAMES = {".git"}
TITLE_RE = re.compile(r"^#\s+(.+?)\s*$", re.MULTILINE)


def should_include(path: Path) -> bool:
    if path.suffix.lower() not in TARGET_EXTENSIONS:
        return False
    if any(part.lower() in IGNORED_DIR_NAMES for part in path.parts):
        return False
    if path.name.lower() in IGNORED_FILE_NAMES:
        return False
    return True


def relative_posix(path: Path, start: Path) -> str:
    return path.relative_to(start).as_posix()


def infer_markdown_title(content: str, fallback: str) -> str:
    match = TITLE_RE.search(content)
    if match:
        return match.group(1).strip()
    return fallback


def output_path_for(raw_file: Path) -> Path:
    parts = raw_file.relative_to(RAW_DIR).parts
    group = parts[0]
    output_root = GROUP_OUTPUTS[group]
    relative_parts = list(parts)
    relative_parts[-1] = f"{relative_parts[-1]}.md"
    return output_root.joinpath(*relative_parts)


def related_wiki_pages(raw_rel: str) -> list[str]:
    pages: list[str] = []
    for page in sorted(WIKI_DIR.rglob("*.md")):
        if not page.is_file():
            continue
        if "原始资料镜像" in page.parts:
            continue
        text = page.read_text(encoding="utf-8")
        if raw_rel in text:
            title_match = re.search(r"^title:\s*(.+?)\s*$", text, re.MULTILINE)
            if title_match:
                pages.append(title_match.group(1).strip().strip('"'))
                continue
            heading_match = re.search(r"^#\s+(.+?)\s*$", text, re.MULTILINE)
            pages.append(heading_match.group(1).strip() if heading_match else page.stem)
    return sorted(set(pages))[:8]


def render_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return ["- 空表"]
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    body = padded[1:] or [[""] * width]
    lines = [
        "| " + " | ".join(cell.replace("|", "\\|") for cell in header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |")
    return lines


def docx_lines(raw_file: Path) -> list[str]:
    document = Document(raw_file)
    lines = ["## 段落内容", ""]
    paragraph_count = 0
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraph_count += 1
        lines.append(f"- {text}")
    if paragraph_count == 0:
        lines.append("- 无非空段落")

    lines.extend(["", "## 表格内容", ""])
    if not document.tables:
        lines.append("- 无表格")
        return lines

    for index, table in enumerate(document.tables, start=1):
        lines.extend([f"### 表格 {index}", ""])
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cell for cell in cells):
                rows.append(cells)
        lines.extend(render_table(rows))
        lines.append("")
    return lines


def xlsx_lines(raw_file: Path) -> list[str]:
    workbook = load_workbook(raw_file, data_only=False, read_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.extend([f"## 工作表：{sheet.title}", ""])
        rows: list[list[str]] = []
        max_width = 0
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            while values and values[-1] == "":
                values.pop()
            if not values:
                continue
            max_width = max(max_width, len(values))
            rows.append(values)
        if not rows:
            lines.append("- 空工作表")
            lines.append("")
            continue
        normalized = [row + [""] * (max_width - len(row)) for row in rows]
        lines.extend(render_table(normalized))
        lines.append("")
    return lines


def render_page(raw_file: Path) -> str:
    raw_rel = relative_posix(raw_file, ROOT)
    related = related_wiki_pages(raw_rel)
    title = f"原始资料-{raw_file.relative_to(RAW_DIR).parts[0]}-{raw_file.stem}"
    tags = ["testing", "raw-source", raw_file.suffix.lower().lstrip(".")]
    summary = f"镜像 `{raw_rel}` 的原始内容，供 Wiki 检索、追溯与后续结构化提炼。"
    lines = [
        "---",
        f"title: {title}",
        "type: archive",
        "status: active",
        "tags:",
    ]
    for tag in tags:
        lines.append(f"  - {tag}")
    lines.extend(
        [
            f"summary: {summary}",
            "source:",
            f"  - {raw_rel}",
            f"updated: {TODAY}",
            "---",
            "",
            f"# {title}",
            "",
            "## 原始文件",
            "",
            f"- 路径：`{raw_rel}`",
            f"- 类型：`{raw_file.suffix.lower()}`",
            "- 说明：本页将 Office 原始文档转成 Wiki 可检索文本，便于测试与知识提炼。",
            "",
            "## 现有 Wiki 关联",
            "",
        ]
    )
    if related:
        lines.extend(f"- [[{item}]]" for item in related)
    else:
        lines.append("- 暂无现有结构化页面直接关联")

    lines.extend(["", "## 原始内容镜像", ""])
    if raw_file.suffix.lower() == ".docx":
        lines.extend(docx_lines(raw_file))
    else:
        lines.extend(xlsx_lines(raw_file))
    lines.append("")
    return "\n".join(lines)


def main() -> None:
    target_groups = {item.strip() for item in sys.argv[1:] if item.strip()}
    count = 0
    for raw_file in sorted(path for path in RAW_DIR.rglob("*") if path.is_file() and should_include(path)):
        group = raw_file.relative_to(RAW_DIR).parts[0]
        if target_groups and group not in target_groups:
            continue
        output_path = output_path_for(raw_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(render_page(raw_file), encoding="utf-8")
        count += 1
    if target_groups:
        group_text = ", ".join(sorted(target_groups))
        print(f"Mirrored {count} Office raw files for groups: {group_text}")
        return
    print(f"Mirrored {count} Office raw files")


if __name__ == "__main__":
    main()
